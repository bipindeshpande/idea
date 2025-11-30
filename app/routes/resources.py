"""Resources routes blueprint - template downloads."""
from flask import Blueprint, request, jsonify, send_file
from typing import Any
import io
import os
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

bp = Blueprint("resources", __name__)


def markdown_to_docx(markdown_content: str, filename: str = "template.docx") -> io.BytesIO:
    """
    Convert markdown content to DOCX format.
    Returns a BytesIO buffer with the DOCX file.
    """
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx library is not installed. Install it with: pip install python-docx")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Add blank line
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            # H1
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            # H2
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            # H3
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            # H4
            p = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s+', line):
            # Bullet or numbered list
            list_text = re.sub(r'^[-*]\s+|\d+\.\s+', '', line)
            p = doc.add_paragraph(list_text, style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            text = line.strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        elif '**' in line:
            # Inline bold
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.strip('*'))
                    run.bold = True
                elif part:
                    p.add_run(part)
        else:
            # Regular paragraph
            # Clean up markdown formatting
            clean_text = line.replace('**', '').replace('*', '').replace('`', '')
            doc.add_paragraph(clean_text)
        
        i += 1
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@bp.post("/api/resources/generate-docx")
def generate_docx():
    """Generate a DOCX file from markdown content."""
    try:
        data = request.get_json()
        content = data.get("content", "")
        filename = data.get("filename", "template.docx")
        
        if not content:
            return jsonify({"success": False, "error": "Content is required"}), 400
        
        # Ensure filename has .docx extension
        if not filename.endswith('.docx'):
            filename = filename.rsplit('.', 1)[0] + '.docx'
        
        # Generate DOCX
        docx_buffer = markdown_to_docx(content, filename)
        
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

